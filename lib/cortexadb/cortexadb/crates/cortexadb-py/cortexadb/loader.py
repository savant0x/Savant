"""
cortexadb.loader — file loading utilities.

Supports loading text from various file formats:
- TXT: Plain text files
- DOCX: Microsoft Word documents
- JSON: JSON files
- MD: Markdown files

Install dependencies:
- TXT: Built-in
- DOCX: pip install cortexadb[docs]
- PDF: pip install cortexadb[pdf]
"""

from __future__ import annotations
from pathlib import Path
from typing import Dict, Any



def load_file(path: str) -> str:
    """
    Load text content from a file.

    Automatically detects format based on extension.

    Args:
        path: Path to the file.

    Returns:
        Extracted text content.

    Raises:
        ValueError: If file format is not supported or not installed.
        FileNotFoundError: If file does not exist.
    """
    p = Path(path)

    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext = p.suffix.lower()

    if ext == ".txt":
        return load_txt(path)
    elif ext == ".md":
        return load_md(path)
    elif ext == ".json":
        return load_json(path)
    elif ext == ".docx":
        return load_docx(path)
    elif ext == ".pdf":
        return load_pdf(path)
    else:
        raise ValueError(
            f"Unsupported file format: {ext}. Supported: .txt, .md, .json, .docx, .pdf"
        )


def load_txt(path: str) -> str:
    """Load plain text file."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_md(path: str) -> str:
    """Load Markdown file (treated as plain text)."""
    return load_txt(path)


def load_json(path: str) -> str:
    """Load JSON file as formatted string."""
    import json

    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, indent=2)


def load_docx(path: str) -> str:
    """Load Microsoft Word document."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError(
            "python-docx not installed. Install with: pip install cortexadb[docs]"
        )

    doc = Document(path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs)


def load_pdf(path: str) -> str:
    """Load PDF document."""
    try:
        import pymupdf
    except ImportError:
        raise ValueError(
            "pymupdf not installed. Install with: pip install cortexadb[pdf]"
        )

    doc = pymupdf.open(path)
    pages = []

    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(text)

    doc.close()
    return "\n\n".join(pages)


def get_file_metadata(path: str) -> Dict[str, Any]:
    """
    Extract metadata from a file.

    Args:
        path: Path to the file.

    Returns:
        Dictionary with file metadata.
    """
    p = Path(path)

    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    stat = p.stat()
    ext = p.suffix.lower()

    metadata = {
        "filename": p.name,
        "file_size": stat.st_size,
        "extension": ext,
    }

    if ext == ".docx":
        try:
            from docx import Document

            doc = Document(path)
            metadata["paragraph_count"] = len(doc.paragraphs)
        except ImportError:
            pass
    elif ext == ".pdf":
        try:
            import pymupdf

            doc = pymupdf.open(path)
            metadata["page_count"] = len(doc)
            doc.close()
        except ImportError:
            pass

    return metadata
